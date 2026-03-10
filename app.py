import streamlit as st
from openai import OpenAI
import docx
from docx import Document
import io
import re

# 1. ПОДКЛЮЧЕНИЕ
API_KEY = st.secrets["OPENROUTER_API_KEY"]

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=API_KEY,
)

# Функция для создания Word-файла из текста ИИ
def create_word_file(content):
    doc = Document()
    doc.add_heading('Протокол разногласий (Евросиб СПб-ТС)', 0)
    doc.add_paragraph(content)
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def read_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    return '\n'.join([para.text for para in doc.paragraphs])

def sanitize_text(text):
    text = re.sub(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', '[EMAIL СКРЫТ]', text)
    text = re.sub(r'\+?[78][-\(]?\d{3}\)?-?\d{3}-?\d{2}-?\d{2}', '[ТЕЛЕФОН СКРЫТ]', text)
    text = re.sub(r'\b\d{20}\b', '[РАСЧЕТНЫЙ СЧЕТ СКРЫТ]', text)
    text = re.sub(r'\b\d{10,15}\b', '[РЕКВИЗИТЫ СКРЫТЫ]', text)
    text = re.sub(r'\b\d{1,3}(?:[.,\s]\d{3})*(?:[.,]\d{1,2})?\s*(?:руб\.|рублей|₽)', '[СУММА СКРЫТА]', text)
    return text

PROFESSIONAL_PROMPT = """Вы — ведущий юрист АО «Евросиб СПб‑ТС». Ваша задача — подготовка протокола разногласий к договору на ремонт вагонов. 
Защищайте интересы Заказчика (качество, сроки, гарантии, неустойки). 
Выдай результат СТРОГО в виде таблицы: Пункт | Редакция Подрядчика | Редакция АО «Евросиб СПб-ТС» | Обоснование."""

st.title("⚖️ Анализатор договоров: Евросиб СПб-ТС")

uploaded_file = st.file_uploader("Загрузите проект договора (.docx)", type=['docx'])
use_censor = st.checkbox("🛡️ Обезличить текст (скрыть суммы, счета, ИНН)", value=True)

if st.button("Сгенерировать протокол разногласий"):
    if uploaded_file is not None:
        raw_text = read_docx(uploaded_file.read())
        final_text = sanitize_text(raw_text) if use_censor else raw_text

        st.info('ИИ анализирует документ... Это может занять до 1 минуты.')
        
        try:
            response = client.chat.completions.create(
                model="openrouter/free", 
                messages=[
                    {"role": "system", "content": PROFESSIONAL_PROMPT},
                    {"role": "user", "content": f"Сделай протокол разногласий:\n\n{final_text}"}
                ]
            )
            
            result_text = response.choices[0].message.content
            st.success("Анализ завершен!")
            st.markdown(result_text)

            # КНОПКА СКАЧИВАНИЯ
            word_file = create_word_file(result_text)
            st.download_button(
                label="📥 Скачать протокол в Word",
                data=word_file,
                file_name="Protocol_Eurosib.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"Ошибка: {e}")


